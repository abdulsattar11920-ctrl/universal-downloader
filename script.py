# script.py
import sys
import os
import re
import json
from yt_dlp import YoutubeDL
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROGRESS_FILE = os.path.join(BASE_DIR, "progress.json")
FINAL_SCRIPT_DIR = r"D:\YouTube_Downloads\scrept"
TEMP_DIR = r"D:\YouTube_Downloads\_temp"

os.makedirs(FINAL_SCRIPT_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)


def write_progress(percent, message, link=""):
    with open(PROGRESS_FILE, "w", encoding="utf-8") as f:
        json.dump({
            "status": "script",
            "percent": percent,
            "message": message,
            "download_url": link
        }, f)


def safe(name):
    return re.sub(r'[\\/:*?"<>|]', '', name)[:180]


url = sys.argv[1]

write_progress(10, "Downloading captions...")

ydl_opts = {
    "skip_download": True,
    "writesubtitles": True,
    "writeautomaticsub": True,
    "subtitleslangs": ["en"],
    "outtmpl": os.path.join(TEMP_DIR, "%(id)s.%(ext)s"),
    "quiet": True,
}

with YoutubeDL(ydl_opts) as ydl:
    info = ydl.extract_info(url, download=True)

title = safe(info.get("title", "video"))

vtt_files = [f for f in os.listdir(TEMP_DIR) if f.endswith(".vtt")]
if not vtt_files:
    write_progress(100, "No captions available")
    sys.exit(1)

vtt = os.path.join(TEMP_DIR, vtt_files[0])

write_progress(60, "Creating document...")

doc = Document()
doc.add_heading(title, level=1)

with open(vtt, "r", encoding="utf-8") as f:
    for line in f:
        if "-->" not in line and line.strip() and not line.strip().isdigit():
            doc.add_paragraph(line.strip())

output = os.path.join(FINAL_SCRIPT_DIR, f"{title}.docx")
doc.save(output)

write_progress(100, "Script ready", f"/get/{title}.docx")
